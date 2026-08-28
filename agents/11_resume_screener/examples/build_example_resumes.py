# ruff: noqa: ISC004
"""Regenerate the binary example resumes (alice.pdf, bob.docx).

Run once, locally, when the resume content needs to change. NOT part
of CI. The output files are committed as binaries.

    uv run --with reportlab python examples/build_example_resumes.py

reportlab is intentionally NOT a workspace dep -- it's only needed
here at author-time. python-docx IS a runtime dep of this agent
already.
"""

from __future__ import annotations

from pathlib import Path

_HERE = Path(__file__).parent
_RESUMES = _HERE / "resumes"


ALICE_CONTENT = [
    ("heading", "Alice Kohler"),
    ("body", "Senior Backend Engineer -- 8 years shipping Python payment infrastructure."),
    ("heading", "Experience"),
    ("subhead", "Bloombox Pay -- Staff Backend Engineer (2022 -- present)"),
    ("body", "Owned the double-entry ledger service handling all merchant payouts. "
             "Rewrote the settlement path from a legacy Django monolith into an "
             "async Python service on top of Postgres and Kafka; sustained 400 TPS "
             "at peak with p99 latency of 80ms."),
    ("body", "Rotated on-call for the payment path from year one (one week every "
             "five). Wrote seven runbooks including the definitive reconciliation "
             "playbook the team still uses. Led three incident reviews."),
    ("body", "Designed the exactly-once event ingestion pipeline: consumer group "
             "with idempotency keys, dead-letter queue with automated replay, and "
             "reconciliation cron that catches any drift within 24 hours."),
    ("subhead", "Firehose Bank -- Senior Backend Engineer (2018 -- 2022)"),
    ("body", "Built the core-banking ledger service in Python + Postgres. "
             "Contributed a small patch to pg_partman for a partitioning edge "
             "case we hit at ~200M rows. Learned formal double-entry accounting "
             "on the job."),
    ("body", "Migrated the deposit-hold service from RabbitMQ to Kafka; handled "
             "the cutover during business hours with zero customer-visible impact "
             "using the outbox pattern."),
    ("subhead", "Reason Systems -- Backend Engineer (2016 -- 2018)"),
    ("body", "First engineering hire at a seed-stage payments startup that later "
             "sold to Firehose. Wore many hats: shipped the first webhook "
             "delivery service, wrote the first Terraform, took the first pages."),
    ("heading", "Skills"),
    ("body", "Python (8y), PostgreSQL, Kafka, async/await, exactly-once patterns, "
             "idempotency, event sourcing, double-entry accounting, distributed "
             "systems (Jepsen-literate), Terraform, pytest, incident response."),
    ("heading", "How I work"),
    ("body", "Fully remote for six years across three continents. Comfortable "
             "with async-first written decisions; strong preference for design "
             "docs over meetings. Have been the on-call primary for a payment "
             "service continuously since 2018."),
]


BOB_CONTENT = [
    ("heading", "Bob Martinez"),
    ("body", "Backend developer, 4 years of Django in production."),
    ("heading", "Experience"),
    ("subhead", "SweetTable Reservations -- Backend Developer (2022 -- present)"),
    ("body", "Own the reservation-booking service written in Django + Django REST "
             "Framework on top of MySQL. Ship features weekly; recent work includes "
             "a waitlist system and SMS reminders."),
    ("body", "Never on-call; our SRE team handles production incidents. I get "
             "paged during business hours if something breaks in my service."),
    ("subhead", "GreenHaus -- Junior Backend Developer (2020 -- 2022)"),
    ("body", "Django + Postgres + Celery for a subscription-box company. Learned "
             "web frameworks and REST API design here."),
    ("heading", "Skills"),
    ("body", "Python (4y, mostly Django + DRF), MySQL, Postgres (basic), Celery, "
             "REST API design, unit testing. Familiar with Redis. Have read about "
             "Kafka but not used it professionally."),
    ("heading", "How I work"),
    ("body", "Hybrid: in the office three days a week, remote the rest. "
             "Comfortable in synchronous meetings; new to async-only written "
             "workflows."),
]


def build_pdf(dest: Path, content: list[tuple[str, str]]) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(dest), pagesize=letter)
    styles = getSampleStyleSheet()
    heading = ParagraphStyle(
        "H", parent=styles["Heading1"], fontSize=14, spaceAfter=6
    )
    subhead = ParagraphStyle(
        "S", parent=styles["Heading2"], fontSize=11, spaceAfter=4
    )
    body = ParagraphStyle("B", parent=styles["BodyText"], fontSize=10, spaceAfter=6)
    flow = []
    for kind, text in content:
        style = {"heading": heading, "subhead": subhead, "body": body}[kind]
        flow.append(Paragraph(text, style))
        flow.append(Spacer(1, 2))
    doc.build(flow)


def build_docx(dest: Path, content: list[tuple[str, str]]) -> None:
    import docx

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    for kind, text in content:
        if kind == "heading":
            doc.add_heading(text, level=1)
        elif kind == "subhead":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    doc.save(str(dest))


def main() -> int:
    build_pdf(_RESUMES / "alice.pdf", ALICE_CONTENT)
    build_docx(_RESUMES / "bob.docx", BOB_CONTENT)
    print(f"wrote {_RESUMES / 'alice.pdf'} and {_RESUMES / 'bob.docx'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
