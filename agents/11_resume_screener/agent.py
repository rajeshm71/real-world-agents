"""Resume screener: agent #11 of real-world-agents.

Technique demonstrated: **batch + ranking with evidence-cross-checked
rationale.** Every prior agent processes one input at a time. #11 is
the first that takes a batch (N resumes) and produces ordered output
(ranked list) with per-item structured scorecards -- the shape a real
HR/recruiting first-pass triage tool actually needs.

Why PydanticAI (same framework as #06 but a different use case):
PydanticAI's `Agent[Deps, Output]` gives us native structured outputs
without us hand-rolling JSON-parse-and-retry (see #02 or #10 for the
hand-rolled shape). Here that matters because the scorecard has four
cross-field validators; letting the framework handle the JSON
contract keeps agent.py focused on batching, ranking, and evidence
grounding.

Real error handling per R5:
- Loader failures (missing file, unsupported suffix, corrupt PDF)
  raise `ScreenerError` at load time with a clear pointer.
- Schema-validation failure on any candidate raises `ScreenerError`
  attaching the raw model output for debugging (matches #10's
  `RetrievalAttempt` shape).
- LLM API failure -> six-branch translator (class-name -> status ->
  message -> generic), matching agents #02-#10 and #13.

Provider stance: OpenAI-only in v1 (mirrors #06 for the same-framework-
same-provider consistency; the README documents the one-line PydanticAI
swap for Anthropic/Gemini).
"""

from __future__ import annotations

import argparse
import functools
import json
import os
import re
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Literal

from pydantic import BaseModel, Field, ValidationError

try:
    from .schemas import (
        DIMENSION_NAMES,
        CandidateScorecard,
        DimensionScore,
        EvidenceExcerpt,
        ResumeInput,
        ScreeningResult,
        _recommendation_for,
    )
except ImportError:
    from schemas import (
        DIMENSION_NAMES,
        CandidateScorecard,
        DimensionScore,
        EvidenceExcerpt,
        ResumeInput,
        ScreeningResult,
        _recommendation_for,
    )

# --- Constants -------------------------------------------------------------

SUPPORTED_PROVIDERS = ("openai", "ollama")
_DEFAULT_PROVIDER = "openai"
_DEFAULT_MODEL_BY_PROVIDER = {
    "openai": "gpt-4o-mini",
    "ollama": "gemma4:e4b",
}
_DEFAULT_MODEL = _DEFAULT_MODEL_BY_PROVIDER["openai"]  # backcompat re-export
_SUPPORTED_SUFFIXES = (".pdf", ".docx", ".md", ".txt")

_PROMPT_PATH = Path(__file__).parent / "prompts" / "system.txt"


# --- Error type ------------------------------------------------------------


@dataclass
class ScreenerAttempt:
    """Partial state attached to ScreenerError. `stage` names where
    things went wrong; the other fields are populated when relevant.

    `completed_scorecards` is populated when a batch fails partway --
    holds the scorecards that succeeded before the offending resume
    hit, so the caller can salvage the completed work instead of
    losing everything on the first failure."""

    stage: Literal["load", "parse", "llm"]
    resume_id: str | None = None
    raw_output: str = ""
    completed_scorecards: list[CandidateScorecard] = field(default_factory=list)


class ScreenerError(Exception):
    """Raised on any user-facing failure: bad path, unsupported format,
    schema-validation failure, LLM API error."""

    def __init__(self, message: str, partial: ScreenerAttempt | None = None):
        super().__init__(message)
        self.message = message
        self.partial = partial


# --- Provider resolution ---------------------------------------------------


def resolve_provider() -> str:
    """LLM_PROVIDER env var, defaulting to openai. Multi-provider swap
    is documented in the README as a one-line PydanticAI change."""
    provider = os.environ.get("LLM_PROVIDER", _DEFAULT_PROVIDER).lower()
    if provider != "mock" and provider not in SUPPORTED_PROVIDERS:
        raise ValueError(
            f"Unknown LLM_PROVIDER: {provider!r}. "
            f"Expected 'mock' or one of {SUPPORTED_PROVIDERS}."
        )
    return provider


@functools.lru_cache(maxsize=1)
def _load_system_prompt() -> str:
    return _PROMPT_PATH.read_text(encoding="utf-8")


# --- Loaders --------------------------------------------------------------


def load_resume(path: str | Path) -> ResumeInput:
    """Load one resume file into a ResumeInput. Dispatches on suffix:
    .pdf -> pypdf, .docx -> python-docx, .md/.txt -> plain read.

    `resume_id` is the lowercased filename stem; `candidate_name` is
    the stem rewritten as a display name (`alice_smith.pdf` ->
    `Alice Smith`)."""
    path = Path(path)
    if not path.exists() or not path.is_file():
        raise ScreenerError(
            f"resume file {str(path)!r} does not exist.",
            partial=ScreenerAttempt(stage="load"),
        )
    suffix = path.suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise ScreenerError(
            f"unsupported resume format {suffix!r} for {path.name!r}; "
            f"supported: {list(_SUPPORTED_SUFFIXES)}.",
            partial=ScreenerAttempt(stage="load"),
        )
    text = _load_resume_text(path)
    if not text.strip():
        raise ScreenerError(
            f"resume {path.name!r} extracted to empty text; the file may "
            "be image-only (OCR not attempted) or corrupt.",
            partial=ScreenerAttempt(stage="load"),
        )
    stem = path.stem
    return ResumeInput(
        resume_id=stem.lower(),
        candidate_name=_stem_to_display_name(stem),
        resume_text=text,
    )


def _load_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf_text(path)
    if suffix == ".docx":
        return _load_docx_text(path)
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ScreenerError(
            "pypdf not installed. Run `uv sync --all-packages` from the "
            "repo root.",
            partial=ScreenerAttempt(stage="load"),
        ) from exc
    try:
        reader = PdfReader(str(path))
        # Strip each page so a trailing newline from extract_text() doesn't
        # compound with the "\n\n" join separator into "\n\n\n\n" runs.
        pages = [(p.extract_text() or "").strip() for p in reader.pages]
    except Exception as exc:
        raise ScreenerError(
            f"pypdf failed on {path.name!r}: {type(exc).__name__}: {exc}.",
            partial=ScreenerAttempt(stage="load"),
        ) from exc
    return "\n\n".join(p for p in pages if p)


def _load_docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ScreenerError(
            "python-docx not installed. Run `uv sync --all-packages` from "
            "the repo root.",
            partial=ScreenerAttempt(stage="load"),
        ) from exc
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        raise ScreenerError(
            f"python-docx failed on {path.name!r}: {type(exc).__name__}: {exc}.",
            partial=ScreenerAttempt(stage="load"),
        ) from exc
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            # A row from a template with all-empty cells prints as
            # " |  | "; skip those so the LLM doesn't see the noise.
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


_STEM_TOKEN_RE = re.compile(r"[_\- ]+")


def _stem_to_display_name(stem: str) -> str:
    """`alice_smith` / `alice-smith` / `alice smith` -> `Alice Smith`."""
    tokens = [t for t in _STEM_TOKEN_RE.split(stem) if t]
    return " ".join(t.capitalize() for t in tokens) or stem


# --- The subset schema the LLM actually returns ---------------------------


class _ScoringOutput(BaseModel):
    """What PydanticAI is told to return. resume_id / candidate_name /
    resume_text are injected by the caller (they're already known),
    so the model doesn't need to echo them back."""

    dimensions: list[DimensionScore] = Field(..., min_length=3, max_length=3)
    overall_score: int = Field(..., ge=0, le=100)
    overall_rationale: str = Field(..., min_length=1)
    recommendation: Literal["strong_yes", "yes", "borderline", "no"]


# --- PydanticAI Agent construction ----------------------------------------


def _build_agent(*, model: str, provider: str = "openai"):
    """Construct the PydanticAI Agent. Lazy-imported so mock-mode paths
    don't require pydantic-ai to be installed. Provider "ollama" routes
    to a local Ollama server via OllamaModel; "openai" keeps the
    existing string-model form."""
    try:
        from pydantic_ai import Agent
    except ImportError as exc:
        raise ScreenerError(
            "pydantic-ai not installed. Run `uv sync --all-packages` from "
            "the repo root."
        ) from exc

    if provider == "ollama":
        from pydantic_ai.models.ollama import OllamaModel
        from pydantic_ai.providers.ollama import OllamaProvider
        from pydantic_ai.settings import ModelSettings

        from common.llm import ollama_base_url
        model_arg: Any = OllamaModel(
            model_name=model,
            provider=OllamaProvider(
                base_url=ollama_base_url(), api_key="ollama"
            ),
        )
        # PydanticAI's OllamaModel treats max_tokens as an overall
        # context cap (prompt + response), not just the response cap
        # -- so give it enough room for JD + resume + system prompt +
        # structured-output overhead + response. 8192 comfortably fits
        # gemma4:e4b's window.
        settings: Any = ModelSettings(max_tokens=8192)
    else:
        model_arg = model
        settings = None

    # Ollama's smaller models produce schema-non-conformant JSON on
    # first try more often than GPT-4o-mini; give the validator a few
    # retries. Openai path keeps the default (1).
    retries = 3 if provider == "ollama" else 1
    return Agent(
        model=model_arg,
        output_type=_ScoringOutput,
        system_prompt=_load_system_prompt(),
        model_settings=settings,
        retries=retries,
    )


# --- Public entry point ----------------------------------------------------


def screen_candidates(
    job_description: str,
    resumes: list[ResumeInput],
    *,
    provider: str | None = None,
    model: str | None = None,
    _agent: Any = None,
) -> ScreeningResult:
    """Score N resumes against a JD, return per-candidate scorecards +
    a ranked-ids permutation.

    Sequential in v1 -- one PydanticAI call per resume -- so the
    reader sees the batching mechanism clearly. `asyncio.gather` on
    the per-candidate calls is the obvious parallel follow-up.
    """
    if not job_description.strip():
        raise ScreenerError("job_description must be non-empty.")
    if not resumes:
        raise ScreenerError("at least one resume is required.")
    ids = [r.resume_id for r in resumes]
    if len(set(ids)) != len(ids):
        seen: set[str] = set()
        dupes: set[str] = set()
        for i in ids:
            if i in seen:
                dupes.add(i)
            seen.add(i)
        raise ScreenerError(
            f"duplicate resume_id(s) not allowed: {sorted(dupes)!r}. "
            "Rename the input files so each stem is unique."
        )

    try:
        resolved = provider or resolve_provider()
    except ValueError as exc:
        raise ScreenerError(str(exc)) from exc

    if resolved == "mock":
        return _mock_screening_result(job_description, resumes)

    resolved_model = model or _DEFAULT_MODEL_BY_PROVIDER.get(
        resolved, _DEFAULT_MODEL
    )
    agent = _agent if _agent is not None else _build_agent(
        model=resolved_model, provider=resolved
    )

    start = time.perf_counter()
    scorecards: list[CandidateScorecard] = []
    for resume in resumes:
        try:
            scorecards.append(
                _score_one(agent, job_description=job_description, resume=resume)
            )
        except ScreenerError as exc:
            # Attach already-completed scorecards so the caller can
            # salvage prior work instead of losing everything on the
            # first parse failure mid-batch.
            if exc.partial is not None:
                exc.partial.completed_scorecards = list(scorecards)
            raise
        except Exception as exc:
            translated = _translate_api_error(exc, resume_id=resume.resume_id)
            if translated.partial is not None:
                translated.partial.completed_scorecards = list(scorecards)
            raise translated from exc

    ranked_ids = _rank(scorecards)
    return ScreeningResult(
        job_description=job_description,
        scorecards=scorecards,
        ranked_ids=ranked_ids,
        run_meta={
            "provider": resolved,
            "model": resolved_model,
            "resume_count": len(resumes),
            "elapsed_seconds": round(time.perf_counter() - start, 3),
        },
    )


def _score_one(
    agent: Any, *, job_description: str, resume: ResumeInput
) -> CandidateScorecard:
    """One PydanticAI call for one resume. Injects the caller-known
    fields (resume_id/name/text) onto the model's scoring output to
    build the full CandidateScorecard, then Pydantic re-validates the
    complete shape (evidence-substring, dimensions-exact-three,
    overall-near-mean, rec-matches-rubric)."""
    user_prompt = _format_user_prompt(job_description, resume)
    result = agent.run_sync(user_prompt)
    scoring: _ScoringOutput = result.output
    try:
        return CandidateScorecard(
            resume_id=resume.resume_id,
            candidate_name=resume.candidate_name,
            resume_text=resume.resume_text,
            dimensions=scoring.dimensions,
            overall_score=scoring.overall_score,
            overall_rationale=scoring.overall_rationale,
            recommendation=scoring.recommendation,
        )
    except ValidationError as exc:
        raise ScreenerError(
            f"scorecard for {resume.resume_id!r} failed cross-field "
            f"validation: {exc}",
            partial=ScreenerAttempt(
                stage="parse",
                resume_id=resume.resume_id,
                raw_output=scoring.model_dump_json(),
            ),
        ) from exc


def _format_user_prompt(job_description: str, resume: ResumeInput) -> str:
    return (
        f"# Job description\n\n{job_description.strip()}\n\n"
        f"# Candidate resume\n"
        f"Candidate: {resume.candidate_name} (id: {resume.resume_id})\n\n"
        f"{resume.resume_text.strip()}"
    )


def _rank(scorecards: list[CandidateScorecard]) -> list[str]:
    """Sort by overall_score desc, resume_id asc (deterministic tie-break)."""
    return [
        s.resume_id
        for s in sorted(scorecards, key=lambda s: (-s.overall_score, s.resume_id))
    ]


# --- Error translation (R5) ------------------------------------------------


def _translate_api_error(
    exc: Exception, *, resume_id: str | None = None
) -> ScreenerError:
    """Turn an openai / pydantic-ai SDK exception into a user-facing
    ScreenerError. Priority: class-name -> status -> message -> generic.
    Same six-branch shape as agents #02-#10 and #13."""
    exc_class_name = type(exc).__name__.lower()
    message_lower = str(exc).lower()
    status = getattr(exc, "status_code", None)
    partial = ScreenerAttempt(stage="llm", resume_id=resume_id)
    if "ratelimiterror" in exc_class_name:
        return _rate_limit_error(partial)
    if "authenticationerror" in exc_class_name or "apikeyerror" in exc_class_name:
        return _auth_error(partial)
    if status == 429:
        return _rate_limit_error(partial)
    if status == 401:
        return _auth_error(partial)
    if "rate limit" in message_lower:
        return _rate_limit_error(partial)
    if "authentication" in message_lower or "api key" in message_lower:
        return _auth_error(partial)
    from common.llm import OLLAMA_CONNECTION_HINT, is_ollama_connection_error
    if is_ollama_connection_error(exc):
        return ScreenerError(OLLAMA_CONNECTION_HINT, partial=partial)
    return ScreenerError(
        f"LLM call failed for resume {resume_id!r}: "
        f"{type(exc).__name__}: {exc}.",
        partial=partial,
    )


def _rate_limit_error(partial: ScreenerAttempt) -> ScreenerError:
    return ScreenerError(
        "OpenAI is rate-limited. Wait a minute and try again.", partial=partial
    )


def _auth_error(partial: ScreenerAttempt) -> ScreenerError:
    return ScreenerError(
        "Authentication failed: check that OPENAI_API_KEY is set. See "
        ".env.example at the repo root.",
        partial=partial,
    )


# --- Mock mode -------------------------------------------------------------


def _mock_screening_result(
    job_description: str, resumes: list[ResumeInput]
) -> ScreeningResult:
    """Scripted per-resume scorecards without importing pydantic-ai or
    openai. Scores derive deterministically from a hash of resume_id
    so different resumes get different scores; evidence is the first
    ~60 chars of the resume text (verbatim by construction). The
    overall_rationale echoes the JD length as an anti-refactor guard
    the tests can assert on."""
    scorecards = [_mock_scorecard(job_description, r) for r in resumes]
    return ScreeningResult(
        job_description=job_description,
        scorecards=scorecards,
        ranked_ids=_rank(scorecards),
        run_meta={
            "provider": "mock",
            "model": "mock",
            "resume_count": len(resumes),
            "elapsed_seconds": 0.0,
        },
    )


def _mock_scorecard(job_description: str, resume: ResumeInput) -> CandidateScorecard:
    seed = sum(ord(c) for c in resume.resume_id)
    scores = {
        "skills_match": 40 + (seed % 51),          # 40..90
        "experience_match": 45 + ((seed // 3) % 46),  # 45..90
        "culture_signal": 50 + ((seed // 7) % 41),    # 50..90
    }
    excerpt_text = _mock_excerpt(resume.resume_text)
    dimensions = [
        DimensionScore(
            name=name,
            score=scores[name],
            rationale=f"[MOCK {name}] scored {scores[name]} for {resume.candidate_name}.",
            evidence=[EvidenceExcerpt(quoted_text=excerpt_text)],
        )
        for name in DIMENSION_NAMES
    ]
    overall = round(sum(scores.values()) / 3)
    return CandidateScorecard(
        resume_id=resume.resume_id,
        candidate_name=resume.candidate_name,
        resume_text=resume.resume_text,
        dimensions=dimensions,
        overall_score=overall,
        overall_rationale=(
            f"[MOCK overall for JD of length {len(job_description)}] "
            f"Averaged three dimensions to {overall} for {resume.candidate_name}."
        ),
        recommendation=_recommendation_for(overall),
    )


def _mock_excerpt(resume_text: str) -> str:
    """First 60 chars of the resume text, trimmed to a token boundary,
    guaranteed to be a verbatim substring of resume_text."""
    snippet = resume_text[:60].rstrip()
    return snippet or resume_text[:1]


# --- CLI ------------------------------------------------------------------


def main() -> int:
    parser = argparse.ArgumentParser(
        prog="resume-screener",
        description=(
            "Screen N resumes against a job description. "
            "Set LLM_PROVIDER=mock for a scripted demo, or supply "
            "OPENAI_API_KEY for real scoring."
        ),
    )
    parser.add_argument(
        "--jd",
        type=Path,
        required=True,
        help="Path to a job description file (any text-readable format).",
    )
    parser.add_argument(
        "--resumes",
        type=Path,
        nargs="+",
        required=True,
        help="One or more resume files (pdf/docx/md/txt).",
    )
    parser.add_argument("--model", type=str, default=None)
    parser.add_argument(
        "--provider", choices=(*SUPPORTED_PROVIDERS, "mock"), default=None
    )
    args = parser.parse_args()

    try:
        jd_text = args.jd.read_text(encoding="utf-8", errors="replace")
        resumes = [load_resume(p) for p in args.resumes]
        result = screen_candidates(
            jd_text,
            resumes,
            provider=args.provider,
            model=args.model,
        )
    except ScreenerError as exc:
        print(f"error: {exc.message}", file=sys.stderr)
        if exc.partial is not None and exc.partial.raw_output:
            print(f"raw model output:\n{exc.partial.raw_output}", file=sys.stderr)
        return 1

    out_path = Path(__file__).parent / "last_run.json"
    out_path.write_text(
        json.dumps(result.model_dump(mode="json"), indent=2), encoding="utf-8"
    )

    print(f"Screened {len(result.scorecards)} candidate(s) against JD "
          f"({len(jd_text)} chars). Ranked:\n")
    by_id = {s.resume_id: s for s in result.scorecards}
    for rank, rid in enumerate(result.ranked_ids, start=1):
        card = by_id[rid]
        top_evidence = next(
            (
                e.quoted_text
                for d in card.dimensions
                for e in d.evidence
            ),
            "(no evidence)",
        )
        print(
            f"  {rank}. [{card.overall_score:>3}] {card.candidate_name} "
            f"({card.recommendation})"
        )
        print(f"       evidence: {top_evidence[:80]!r}")
    print(f"\nFull result written to {out_path}.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
